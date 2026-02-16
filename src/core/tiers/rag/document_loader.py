"""Load and process documents for RAG"""

from pathlib import Path
from typing import List, Dict

try:
    import pypdf
except ImportError:
    pypdf = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None


class Document:
    """Simple document class"""
    def __init__(self, text: str, metadata: Dict):
        self.text = text
        self.metadata = metadata


class DocumentLoader:
    """Load documents from various formats"""

    def __init__(self, source_dir: str = "data/documents"):
        self.source_dir = Path(source_dir)

    def load_all(self) -> List[Document]:
        """Load all documents from source directory"""
        documents = []

        if not self.source_dir.exists():
            print(f"WARNING: Document directory not found: {self.source_dir}")
            print("Creating directory...")
            self.source_dir.mkdir(parents=True, exist_ok=True)
            return documents

        # Find all supported files
        for file_path in self.source_dir.rglob("*"):
            if file_path.is_file():
                try:
                    doc = self._load_file(file_path)
                    if doc:
                        documents.append(doc)
                except Exception as e:
                    print(f"WARNING: Error loading {file_path.name}: {e}")

        print(f"Loaded {len(documents)} documents")
        return documents

    def _load_file(self, file_path: Path) -> Document:
        """Load single file based on extension"""
        ext = file_path.suffix.lower()

        if ext == ".pdf":
            return self._load_pdf(file_path)
        if ext == ".docx":
            return self._load_docx(file_path)
        if ext in [".txt", ".md"]:
            return self._load_text(file_path)
        return None

    def _load_pdf(self, file_path: Path) -> Document:
        """Load PDF file"""
        if pypdf is None:
            raise RuntimeError("pypdf is not installed (pip install pypdf)")

        text = ""
        with open(file_path, "rb") as f:
            pdf = pypdf.PdfReader(f)
            for page in pdf.pages:
                text += page.extract_text() + "\n"

        return Document(
            text=text.strip(),
            metadata={
                "source": str(file_path),
                "filename": file_path.name,
                "type": "pdf",
            },
        )

    def _load_docx(self, file_path: Path) -> Document:
        """Load DOCX file"""
        if DocxDocument is None:
            raise RuntimeError("python-docx is not installed (pip install python-docx)")

        doc = DocxDocument(file_path)
        text = "\n".join([para.text for para in doc.paragraphs])

        return Document(
            text=text.strip(),
            metadata={
                "source": str(file_path),
                "filename": file_path.name,
                "type": "docx",
            },
        )

    def _load_text(self, file_path: Path) -> Document:
        """Load text file"""
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()

        return Document(
            text=text.strip(),
            metadata={
                "source": str(file_path),
                "filename": file_path.name,
                "type": "text",
            },
        )


__all__ = ["DocumentLoader", "Document"]
